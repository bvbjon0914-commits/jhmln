"""
Generiert die Word-Vorlagen für alle Auskunftsarten.

Layout orientiert sich an den Vonovia-Referenzvorlagen
(Serienbrief_Bauamt_Vorlage.docx / Serienbrief_Grundbuchamt_Vorlage.docx):
Vonovia-Absenderzeile, Trennlinie, Empfängerblock, rechtsbündiges Datum,
fett gesetzter Betreff mit Objektadresse und interner Referenz, Fließtext
mit auskunftsart-spezifischer Rechtsgrundlage, "Betroffenes Objekt/
Grundstück"-Block, Absatz zum berechtigten Interesse, Checkbox-Optionen,
Gebührenhinweis, DSGVO-Hinweis (kursiv/grau), Grußformel, Signatur-
Platzhalter, Anlagenzeile und optionaler Impressum-Fußblock.

Wird einmalig ausgeführt, um die Vorlagen unter /templates zu erzeugen.
Die {{ platzhalter }} bleiben erhalten und werden von docxtpl beim
tatsächlichen Erzeugen der Anschreiben (DocumentGenerationService) befüllt.
"""

import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")

GRAY = RGBColor(0x59, 0x59, 0x59)
CHECKBOX = "☐"  # ☐


# ---------------------------------------------------------------------------
# Auskunftsart-spezifische Inhalte
#
# subject_line1 / subject_line2: fett gesetzter zweizeiliger Betreff
# legal_basis:   Rechtsgrundlage im Fließtext (Bundesland wird eingesetzt)
# object_label:  "Objekt" oder "Grundstück"
# interest_hint: Klammerbeispiel für das berechtigte Interesse
# checkboxes:    Liste der Anfrageoptionen
# extra_note:    optionaler zusätzlicher Hinweisabsatz (kursiv, wie im
#                Grundbuch-Muster die Anmerkung zu Gemarkung/Flur)
# ---------------------------------------------------------------------------
REQUEST_TYPE_TEXTS = {
    "grundbuch": {
        "subject_line1": "Antrag auf Grundbucheinsicht / Erteilung eines Grundbuchauszugs – Objekt",
        "legal_basis": (
            "hiermit beantrage(n) ich/wir gemäß § 12 der Grundbuchordnung (GBO) "
            "Einsicht in das Grundbuch bzw. die Erteilung eines Grundbuchauszugs für das "
            "nachfolgend bezeichnete Grundstück."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Erbfolge, Bevollmächtigung durch den Eigentümer, "
            "Verkehrswertermittlung im Auftrag"
        ),
        "checkboxes": ["einfachen Grundbuchauszugs", "beglaubigten Grundbuchauszugs"],
        "intro": "Wir bitten um Übersendung eines:",
        "extra_note": (
            "(Gemarkung, Flur und Flurstück sind uns derzeit nicht bekannt – wir bitten "
            "um Identifikation anhand der oben genannten Adresse.)"
        ),
    },
    "bauakten": {
        "subject_line1": "Antrag auf Akteneinsicht / Bauaktenauskunft – Objekt",
        "legal_basis": (
            "hiermit beantrage(n) ich/wir Akteneinsicht bzw. Auskunft aus der Bauakte für "
            "das nachfolgend bezeichnete Objekt. Die Anfrage stützt sich auf die "
            "einschlägigen Vorschriften der Bauordnung des Landes {state} sowie – so "
            "weit einschlägig – auf § 29 des Verwaltungsverfahrensgesetzes (VwVfG) bzw. "
            "das entsprechende Landesverwaltungsverfahrensgesetz."
        ),
        "object_label": "Objekt",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "gutachterliche Tätigkeit"
        ),
        "checkboxes": [
            "Übersendung von Kopien der vorhandenen Bauakte(n) (insbesondere Baugenehmigung, genehmigte Baupläne/Lagepläne, Abnahmeprotokolle)",
            "Mitteilung, ob und in welchem Umfang Bauakten zu diesem Objekt bei Ihnen vorliegen",
            "Terminvereinbarung zur Einsichtnahme vor Ort in Ihren Diensträumen",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "baulasten": {
        "subject_line1": "Antrag auf Auskunft aus dem Baulastenverzeichnis – Objekt",
        "legal_basis": (
            "hiermit beantrage(n) ich/wir eine Auskunft aus dem Baulastenverzeichnis nach "
            "der Landesbauordnung des Landes {state} für das nachfolgend bezeichnete "
            "Grundstück."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vorbereitung einer Bebauung/Veräußerung"
        ),
        "checkboxes": [
            "Übersendung eines aktuellen Auszugs aus dem Baulastenverzeichnis",
            "Mitteilung, ob und welche Baulasten zu diesem Grundstück eingetragen sind",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "altlasten": {
        "subject_line1": "Antrag auf Altlastenauskunft – Objekt",
        "legal_basis": (
            "hiermit beantrage(n) ich/wir eine Auskunft aus dem Altlasten-/Bodenschutzkataster "
            "gemäß § 21 des Bundes-Bodenschutzgesetzes (BBodSchG) i. V. m. dem "
            "Bodenschutzrecht des Landes {state} für das nachfolgend bezeichnete Grundstück."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vorbereitung einer Transaktion/Bebauung"
        ),
        "checkboxes": [
            "Mitteilung, ob das Grundstück im Altlasten-/Bodenschutzkataster erfasst ist",
            "Übersendung vorhandener Unterlagen (Gutachten, Untersuchungsberichte)",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "erschliessung": {
        "subject_line1": "Anfrage zu Erschließungsbeiträgen / Anliegerbescheinigung – Objekt",
        "legal_basis": (
            "hiermit bitte(n) ich/wir um Auskunft über den Stand der Erschließungsbeiträge "
            "gemäß §§ 127 ff. des Baugesetzbuches (BauGB) sowie um Ausstellung einer "
            "Anliegerbescheinigung für das nachfolgend bezeichnete Grundstück."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vorbereitung einer Transaktion"
        ),
        "checkboxes": [
            "Mitteilung, ob Erschließungsbeiträge noch offen sind oder abgerechnet wurden",
            "Ausstellung einer Anliegerbescheinigung (Unbedenklichkeitsbescheinigung)",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "denkmalschutz": {
        "subject_line1": "Anfrage zum Denkmalschutzstatus – Objekt",
        "legal_basis": (
            "hiermit bitte(n) ich/wir um Auskunft, ob für das nachfolgend bezeichnete Objekt "
            "eine Eintragung in die Denkmalliste nach dem Denkmalschutzgesetz des Landes "
            "{state} besteht."
        ),
        "object_label": "Objekt",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vorbereitung von Umbau-/Modernisierungsmaßnahmen"
        ),
        "checkboxes": [
            "Mitteilung, ob das Objekt in die Denkmalliste eingetragen ist",
            "Übersendung eines Auszugs aus der Denkmalliste bzw. relevanter Unterlagen",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "bodendenkmalschutz": {
        "subject_line1": "Anfrage zum Bodendenkmalschutz – Objekt",
        "legal_basis": (
            "hiermit bitte(n) ich/wir um Auskunft, ob für das nachfolgend bezeichnete "
            "Grundstück Bodendenkmäler nach dem Denkmalschutzgesetz des Landes {state} "
            "bekannt oder eingetragen sind."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vorbereitung von Baumaßnahmen"
        ),
        "checkboxes": [
            "Mitteilung, ob Bodendenkmäler auf dem Grundstück bekannt/eingetragen sind",
            "Übersendung relevanter Unterlagen bzw. Auflagen für Baumaßnahmen",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "wasserschutz": {
        "subject_line1": "Anfrage zum Wasserschutzgebiet – Objekt",
        "legal_basis": (
            "hiermit bitte(n) ich/wir um Auskunft, ob sich das nachfolgend bezeichnete "
            "Grundstück innerhalb eines festgesetzten Wasserschutzgebietes gemäß § 51 des "
            "Wasserhaushaltsgesetzes (WHG) i. V. m. dem Wassergesetz des Landes {state} "
            "befindet."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vorbereitung einer Transaktion/Bebauung"
        ),
        "checkboxes": [
            "Mitteilung, ob das Grundstück in einem Wasserschutzgebiet liegt (inkl. Zone)",
            "Übersendung geltender Schutzgebietsverordnungen/Auflagen",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "hochwasserschutz": {
        "subject_line1": "Anfrage zum Hochwasserrisiko – Objekt",
        "legal_basis": (
            "hiermit bitte(n) ich/wir um Auskunft, ob das nachfolgend bezeichnete Grundstück "
            "in einem festgesetzten Überschwemmungsgebiet gemäß § 76 des "
            "Wasserhaushaltsgesetzes (WHG) i. V. m. dem Wassergesetz des Landes {state} liegt."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Risikobewertung im Rahmen der Bestandsverwaltung"
        ),
        "checkboxes": [
            "Mitteilung, ob das Grundstück in einem Überschwemmungsgebiet liegt",
            "Übersendung von Hochwassergefahrenkarten bzw. relevanter Unterlagen",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "kampfmittel": {
        "subject_line1": "Anfrage zur Kampfmittelfreiheit – Objekt",
        "legal_basis": (
            "hiermit bitte(n) ich/wir um Auskunft über eine mögliche Kampfmittelbelastung des "
            "nachfolgend bezeichneten Grundstücks gemäß den einschlägigen Vorschriften zur "
            "Kampfmittelräumung des Landes {state} sowie ggf. um Mitteilung des weiteren "
            "Vorgehens."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vorbereitung von Baumaßnahmen"
        ),
        "checkboxes": [
            "Auswertung der Luftbilder / Mitteilung des Ergebnisses",
            "Mitteilung, ob eine Sondierung/Räumung erforderlich ist",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
    "kataster": {
        "subject_line1": "Antrag auf Auskunft aus dem Liegenschaftskataster – Objekt",
        "legal_basis": (
            "hiermit beantrage(n) ich/wir eine aktuelle Auskunft aus dem Liegenschaftskataster "
            "(Flurstücksnachweis / Auszug aus der Liegenschaftskarte) nach dem Vermessungs- "
            "und Katastergesetz des Landes {state} für das nachfolgend bezeichnete "
            "Grundstück."
        ),
        "object_label": "Grundstück",
        "interest_hint": (
            "Eigentümerstellung, Rechtsnachfolge, Bevollmächtigung durch den Eigentümer, "
            "Vermessungs-/Vermarktungszwecke"
        ),
        "checkboxes": [
            "Übersendung eines aktuellen Flurstücksnachweises",
            "Übersendung eines Auszugs aus der Liegenschaftskarte",
        ],
        "intro": "Wir bitten höflich um:",
        "extra_note": None,
    },
}


def _set_bottom_border(paragraph, color="999999", size=6):
    """Fügt einem Absatz eine untere Rahmenlinie hinzu (Trennlinie)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _gray_small(paragraph, text, size=8, italic=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GRAY
    run.italic = italic
    return run


def build_template(code: str, texts: dict) -> str:
    """Erstellt ein einzelnes DOCX-Template und gibt den Dateipfad zurück."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.05

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # --- Absenderblock (Vonovia SE) -----------------------------------
    p = doc.add_paragraph()
    run = p.add_run("Vonovia SE")
    run.bold = True

    p = doc.add_paragraph()
    _gray_small(p, "Universitätsstr. 133 · 44803 Bochum · Tel. [Telefon] · [E-Mail-Adresse]")

    p = doc.add_paragraph()
    _gray_small(p, "Vonovia SE · Universitätsstr. 133 · 44803 Bochum", size=7.5)
    _set_bottom_border(p)

    doc.add_paragraph()

    # --- Empfängerblock ------------------------------------------------
    doc.add_paragraph("{{ authority_name }}")
    doc.add_paragraph("{{ authority_department }}")
    doc.add_paragraph("{{ authority_street }} {{ authority_house_number }}")
    doc.add_paragraph("{{ authority_postal_code }} {{ authority_city }}")

    doc.add_paragraph()

    # --- Datum, rechtsbündig -------------------------------------------
    date_p = doc.add_paragraph("Bochum, den {{ current_date }}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # --- Betreff (zweizeilig, fett) --------------------------------------
    subj1 = doc.add_paragraph()
    r = subj1.add_run(texts["subject_line1"] + " ")
    r.bold = True
    r2 = subj1.add_run("{{ building_street }} {{ building_house_number }}")
    r2.bold = True

    subj2 = doc.add_paragraph()
    r = subj2.add_run(
        "({{ building_postal_code }} {{ building_city }}  –  interne Referenz "
        "{{ internal_reference }})"
    )
    r.bold = True

    doc.add_paragraph()

    # --- Anrede + Anfragetext --------------------------------------------
    doc.add_paragraph("Sehr geehrte Damen und Herren,")
    doc.add_paragraph()

    body_text = texts["legal_basis"].format(state="{{ building_state }}")
    doc.add_paragraph(body_text)
    doc.add_paragraph()

    # --- Objektangaben -----------------------------------------------------
    label_p = doc.add_paragraph()
    label_p.add_run(f"Betroffenes {texts['object_label']}:").bold = True
    doc.add_paragraph("Gemeinde: {{ building_city }}")
    doc.add_paragraph(
        "Adresse: {{ building_street }} {{ building_house_number }}, "
        "{{ building_postal_code }} {{ building_city }}"
    )
    doc.add_paragraph("Bundesland: {{ building_state }}")
    doc.add_paragraph("Interne Referenz-Nr.: {{ internal_reference }}")

    if texts.get("extra_note"):
        note_p = doc.add_paragraph()
        _gray_small(note_p, texts["extra_note"], size=9.5, italic=True)

    doc.add_paragraph()

    # --- Berechtigtes Interesse ---------------------------------------------
    interest_label = "erforderliche berechtigte Interesse"
    doc.add_paragraph(
        f"Das für die Anfrage {interest_label} ergibt sich aus [BITTE ERGÄNZEN: z. B. "
        f"{texts['interest_hint']}]. Einen entsprechenden Nachweis (z. B. Grundbuchauszug, "
        "Vollmacht, Auftragsschreiben) fügen wir diesem Schreiben als Anlage bei."
    )
    doc.add_paragraph()

    # --- Checkboxen ---------------------------------------------------------
    doc.add_paragraph(texts["intro"])
    for option in texts["checkboxes"]:
        doc.add_paragraph(f"{CHECKBOX} {option}")
    doc.add_paragraph()

    # --- Gebührenhinweis -----------------------------------------------
    doc.add_paragraph(
        "Für die Bearbeitung dieser Anfrage anfallende Verwaltungsgebühren gemäß der "
        "für Sie geltenden Gebührenordnung werden von uns übernommen. Sollten die Kosten "
        "voraussichtlich [BETRAG] EUR übersteigen, bitten wir um vorherige Mitteilung."
    )
    doc.add_paragraph()

    # --- DSGVO-Hinweis (kursiv, grau) ----------------------------------
    dsgvo_p = doc.add_paragraph()
    _gray_small(
        dsgvo_p,
        "Datenschutzhinweis: Die von uns im Rahmen dieser Anfrage übermittelten sowie die "
        "im Zuge der Bearbeitung ggf. erhaltenen personenbezogenen Daten werden von uns "
        "ausschließlich zur Bearbeitung dieses Anliegens auf Grundlage von Art. 6 Abs. 1 "
        "lit. f DSGVO (berechtigtes Interesse) bzw. Art. 6 Abs. 1 lit. c DSGVO (rechtliche "
        "Verpflichtung) verarbeitet und nicht an unbefugte Dritte weitergegeben.",
        size=10,
        italic=True,
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Für Rückfragen stehen wir Ihnen gerne unter den oben genannten Kontaktdaten zur "
        "Verfügung. Vielen Dank für Ihre Unterstützung."
    )
    doc.add_paragraph()

    doc.add_paragraph("Mit freundlichen Grüßen")
    doc.add_paragraph()
    doc.add_paragraph("[Name]")
    doc.add_paragraph("[Funktion / Firma]")
    doc.add_paragraph()

    anlagen_p = doc.add_paragraph()
    r = anlagen_p.add_run("Anlagen: [ggf. Vollmacht / Eigentumsnachweis / Auftragsschreiben]")
    r.italic = True

    # --- Impressum-Fußblock -------------------------------------------
    doc.add_paragraph()
    impressum_p = doc.add_paragraph()
    _set_bottom_border(impressum_p, color="CCCCCC", size=4)
    impressum_p.add_run(
        "Vonovia SE · Sitz: [Ort] · Registergericht: [Amtsgericht] · "
        "Registernummer: [HRB ...] · Vorstand: [Name(n)] · USt-IdNr.: [DE ...]"
    ).font.size = Pt(8)
    hint_p = doc.add_paragraph()
    _gray_small(
        hint_p,
        "(Hinweis: Dieser Impressum-Block ist nur erforderlich, wenn der Absender ein im "
        "Handelsregister eingetragenes Unternehmen ist, § 37a HGB / § 35a AktG. Bitte die "
        "eckigen Klammern vor Versand ausfüllen.)",
        size=7.5,
        italic=True,
    )

    output_path = os.path.join(TEMPLATES_DIR, f"{code}.docx")
    doc.save(output_path)
    return output_path


def main():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    for code, texts in REQUEST_TYPE_TEXTS.items():
        path = build_template(code, texts)
        print(f"✓ Erstellt: {path}")


if __name__ == "__main__":
    main()
